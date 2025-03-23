from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from datetime import datetime
import zipfile

def crear_documento_sponsorship_of_event(dataframe):
    # Crear un nuevo documento
    documento = Document()

    # Agregar el título con estilo y color personalizado
    titulo = documento.add_paragraph()
    run_titulo = titulo.add_run('Sponsorship of Event')
    run_titulo.font.size = Pt(16)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(128, 0, 128)  # Color morado
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar secciones con encabezados estilizados
    def agregar_encabezado(texto):
        parrafo = documento.add_paragraph()
        run = parrafo.add_run(texto)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Color azul
        parrafo.space_after = Pt(6)

    def agregar_bullet_point(campo, valor):
        parrafo = documento.add_paragraph(style='List Bullet')
        run_campo = parrafo.add_run(f"{campo}: ")
        run_campo.font.size = Pt(11)
        run_campo.font.bold = True
        run_valor = parrafo.add_run(f"{valor}")
        run_valor.font.size = Pt(11)

    # Leer los datos del DataFrame
    datos = dataframe.iloc[0].to_dict()

    # Detalles del evento
    agregar_encabezado("Detalles del Evento")
    agregar_bullet_point("Nombre del evento", f"Sponsorship of Event/Activity {datos.get('event_name', '')}")
    agregar_bullet_point("Owner", datos.get('owner', ''))
    if datos.get("delegate", "") != "":
        agregar_bullet_point("Delegate", datos.get('delegate', ''))
    else:
        agregar_bullet_point("Delegate", "N/A")
    agregar_bullet_point("Fecha de inicio", datos.get("start_date", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Fecha de fin", datos.get("end_date", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Tipo de evento", datos.get("event_type", ""))
    if datos.get("venue", "") == "":
        agregar_bullet_point("Sede", "N/A")
    else:
        agregar_bullet_point("Sede", datos.get("venue", "N/A"))
    if datos.get("city", "") == "":
        agregar_bullet_point("Ciudad","N/A")
    else:
        agregar_bullet_point("Ciudad", datos.get("city", "N/A"))
    agregar_bullet_point("Número de asistentes", datos.get("num_attendees", 0))
    agregar_bullet_point("Perfil de asistentes", datos.get("attendee_profile", ""))
    agregar_bullet_point("Objetivo del evento", datos.get("event_objetive", ""))

    # Detalles del firmante
    agregar_encabezado("Detalles del Organizador")
    agregar_bullet_point("Nombre de la organización", datos.get("organization_name", ""))
    agregar_bullet_point("CIF", datos.get("organization_cif", ""))
    agregar_bullet_point("Nombre y apellidos del firmante", f"{datos.get('signer_first_name', '')}")
    agregar_bullet_point("Cargo del firmante", datos.get("signer_position", ""))
    agregar_bullet_point("Email del firmante", datos.get("signer_email", ""))

    # Detalles del patrocinio
    agregar_encabezado("Detalles del Patrocinio")
    agregar_bullet_point("Nombre del evento", f"Sponsorship of Event/Activity {datos.get('event_name', '')}")
    agregar_bullet_point("Importe", f"{datos.get('amount', 0.0)} €")
    agregar_bullet_point("Tipo de pago", datos.get("payment_type", ""))
    if datos.get("payment_type") == "Pago a través de la secretaría técnica (ST)":
        agregar_bullet_point("Nombre ST", datos.get("name_st", ""))
    if datos.get("associated_product") == "":
        agregar_bullet_point("Producto asociado", "N/A")
    else:
        agregar_bullet_point("Producto asociado", datos.get("associated_product", "N/A"))
    agregar_bullet_point("Descripción del evento", datos.get("short_description", ""))
    agregar_bullet_point("Contraprestaciones", datos.get("benefits", ""))
    agregar_bullet_point("Patrocinador único o mayoritario", datos.get("exclusive_sponsorship", "No"))
    agregar_bullet_point("Patrocinio recurrente", datos.get("recurrent_sponsorship", ""))
    if datos.get("recurrent_sponsorship", "No") == "Sí":
        agregar_bullet_point("Detalles del patrocinio recurrente", datos.get("recurrent_text", ""))


    nombre_zip = f"Sponshorship_Event {datos.get('event_name', '')}.zip"
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    archivo_zip = os.path.join(output_dir, nombre_zip)

    with zipfile.ZipFile(archivo_zip, 'w') as zipf:
        nombre_archivo = 'Sponsorship_Event_Plantilla.docx'
        archivo_docx = os.path.join(output_dir, nombre_archivo)
        documento.save(archivo_docx)
        zipf.write(archivo_docx, os.path.basename(archivo_docx))

        doc1 = datos.get("documentosubido_1_event", None)
        doc2 = datos.get("documentosubido_2_event", None)
        doc3 = datos.get("documentosubido_3_event", None)
        doc4 = datos.get("documentosubido_4_event", None)
        doc5 = datos.get("documentosubido_5_event", None)

        if doc1 is not None:
            #print(str(doc1.type))
            file_type = str(doc1.name).split(".")[-1]
            doc1_path = os.path.join(output_dir, f"AgendaEvento.{file_type}")  
            with open(doc1_path, "wb") as f:
                f.write(doc1.getbuffer()) 
            zipf.write(doc1_path, os.path.basename(doc1_path))  

        if doc2 is not None:
            file_type = str(doc2.name).split(".")[-1]
            doc2_path = os.path.join(output_dir, f"SolicitudPatrocinio.{file_type}")  
            with open(doc2_path, "wb") as f:
                f.write(doc2.getbuffer()) 
            zipf.write(doc2_path, os.path.basename(doc2_path))  

        if doc3 is not None and doc3 != "":
            file_type = str(doc3.name).split(".")[-1]
            doc3_path = os.path.join(output_dir, f"DossierComercial.{file_type}")  
            with open(doc3_path, "wb") as f:
                f.write(doc3.getbuffer()) 
            zipf.write(doc3_path, os.path.basename(doc3_path))  
        
        if doc4 is not None:
            contador = 1
            for uploaded_file in doc4:
                file_type = str(uploaded_file.name).split(".")[-1]
                doc4_path = os.path.join(output_dir,  f"DocumentoAdicional_{contador}.{file_type}") 
                with open(doc4_path, "wb") as f:
                    f.write(uploaded_file.getbuffer()) 
                zipf.write(doc4_path, os.path.basename(doc4_path))  
                contador += 1

        if doc5 is not None and doc5 != "":
            file_type = str(doc5.name).split(".")[-1]
            doc5_path = os.path.join(output_dir, f"DossierComercial.{file_type}")  
            with open(doc5_path, "wb") as f:
                f.write(doc5.getbuffer()) 
            zipf.write(doc5_path, os.path.basename(doc5_path)) 

    os.remove(archivo_docx)
    if doc1 is not None:
        os.remove(doc1_path)
    if doc2 is not None:
        os.remove(doc2_path)
    if doc3 is not None and doc3 != "":
        os.remove(doc3_path)
    if doc4 is not None:
        os.remove(doc4_path)
    if doc5 is not None and doc5 != "":
        os.remove(doc5_path)
    print(f'Documento y archivos añadidos al ZIP: {nombre_zip}')

    return documento, archivo_zip



def crear_documento_advisory(data):
    documento = Document()

    # Agregar el título
    titulo = documento.add_paragraph()
    run_titulo = titulo.add_run('Advisory Board')
    run_titulo.font.size = Pt(16)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def agregar_encabezado(texto):
        parrafo = documento.add_paragraph()
        run = parrafo.add_run(texto)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Azul
        parrafo.space_after = Pt(6)

    def agregar_bullet_point(campo, valor):
        parrafo = documento.add_paragraph(style='List Bullet')
        run_campo = parrafo.add_run(f"{campo}: ")
        run_campo.font.size = Pt(11)
        run_campo.font.bold = True
        run_valor = parrafo.add_run(f"{valor}")
        run_valor.font.size = Pt(11)

    # Agregar secciones
    agregar_encabezado("Detalles del Evento")
    agregar_bullet_point("Nombre", data.get('nombre_evento_ab', ''))
    agregar_bullet_point("Owner", data.get('owner_ab', ''))
    if data.get("delegate_ab") != "":
        agregar_bullet_point("Delegate", data.get('delegate_ab', ""))
    else:
        agregar_bullet_point("Delegate", "N/A")
    #agregar_bullet_point("Descripción y objetivo", data.get("descripcion_objetivo_ab", ""))
    agregar_bullet_point("Fecha de inicio", data.get("start_date_ab", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Fecha de fin", data.get("end_date_ab", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Tipo de evento", data.get("tipo_evento_ab", ""))
    if data.get("sede_ab", "") == "":
        agregar_bullet_point("Sede", "N/A")
    else:
        agregar_bullet_point("Sede", data.get("sede_ab", "N/A"))
    if data.get("ciudad_ab", "") == "":
        agregar_bullet_point("Ciudad", "N/A")
    else:
        agregar_bullet_point("Ciudad", data.get("ciudad_ab", "N/A"))
    #agregar_bullet_point("Público objetivo", data.get("publico_objetivo_ab", ""))
    
    agregar_encabezado("Detalles de la Actividad")
    if data.get("producto_asociado_ab", "") != "":
        agregar_bullet_point("Producto asociado", data.get("producto_asociado_ab", ""))
    else:
        agregar_bullet_point("Producto asociado", "N/A")
    agregar_bullet_point("Estado de la aprobación", data.get("estado_aprobacion_ab", ""))
    agregar_bullet_point("Descripción del servicio", data.get("descripcion_servicio_ab", ""))
    agregar_bullet_point("Necesidad de la reunión y resultados esperados", data.get("necesidad_reunion_ab", ""))
    agregar_bullet_point("Actividad en el departamento en los últimos 12 meses", data.get("otra_actividad_departamento_ab", ""))
    agregar_bullet_point("Actividad en otro departamento en los últimos 12 meses", data.get("otra_actividad_otro_departamento_ab", ""))

    agregar_encabezado("Detalle nº Participantes")
    agregar_bullet_point("Nº de participantes", data.get("num_participantes_ab", ""))
    agregar_bullet_point("Justificación de número de participantes", data.get("justificacion_participantes_ab", ""))
    criterios = ", ".join(data.get("criterios_seleccion_ab", []))
    agregar_bullet_point("Criterios de selección", criterios)

    agregar_encabezado("Logística de los Participantes")
    agregar_bullet_point("Desplazamiento de participantes", data.get("desplazamiento_ab", ""))
    agregar_bullet_point("Alojamiento de participantes", data.get("alojamiento_ab", ""))
    if data.get("alojamiento_ab") == "Sí":
        agregar_bullet_point("Nº de noches", data.get("num_noches_ab", ""))
        agregar_bullet_point("Hotel", data.get("hotel_ab", ""))
    
    agregar_encabezado("Detalles de los Participantes del Advisory")
    def agregar_tabla_participantes(participantes):
        tabla = documento.add_table(rows=1, cols=8)
        tabla.style = 'Table Grid'  # Aplicar bordes a la tabla
        encabezados = ["Nombre", "DNI", "Tier", "Centro de trabajo", "Email", "Cobra a través de sociedad", "Honorarios", "Tiempos"]
        hdr_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            hdr_cells[i].text = encabezado
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i]._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w')))
        
        for participante in participantes.values():
            id_participante = participante["id"]
            row_cells = tabla.add_row().cells
            #row_cells[0].text = participante.get(f"nombre_{id_participante}", "").split('-')[0]
            row_cells[0].text = participante.get(f"nombre_{id_participante}", "").get("result", "").rsplit('-', 1)[0] 
            row_cells[1].text = participante.get(f"dni_{id_participante}", "")
            row_cells[2].text = participante.get(f"tier_{id_participante}", "")
            row_cells[3].text = participante.get(f"centro_trabajo_{id_participante}", "")
            row_cells[4].text = participante.get(f"email_{id_participante}", "")
            if participante.get(f"cobra_sociedad_{id_participante}", "") == "Sí":
                row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") + ", con sociedad: " + participante.get(f"nombre_sociedad_{id_participante}", "")
            else:
                row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "")
            row_cells[6].text = str(participante.get(f"honorarios_{id_participante}", ""))
            row_cells[7].text = f"Preparación: {participante.get(f'preparacion_horas_{id_participante}', '')} horas y {participante.get(f'preparacion_minutos_{id_participante}', '')} minutos, Ponencia: {participante.get(f'ponencia_horas_{id_participante}', '')} horas y {participante.get(f'ponencia_minutos_{id_participante}', '')} minutos"
            
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w')))
    
    agregar_tabla_participantes(data.get("participantes_ab", {}))

    nombre_zip = f"Advisory_Board {data.get('nombre_evento_ab', '')}.zip"
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    archivo_zip = os.path.join(output_dir, nombre_zip)

    with zipfile.ZipFile(archivo_zip, 'w') as zipf:
        nombre_archivo = 'Advisory_Board_Plantilla.docx'
        archivo_docx = os.path.join(output_dir, nombre_archivo)
        documento.save(archivo_docx)
        zipf.write(archivo_docx, os.path.basename(archivo_docx))

        doc1 = data.get("documentosubido_1", None)
        doc2 = data.get("documentosubido_2", None)

        if doc1 is not None:
            file_type = str(doc1.name).split(".")[-1]
            doc1_path = os.path.join(output_dir, f"ProgramaEvento.{file_type}")  
            with open(doc1_path, "wb") as f:
                f.write(doc1.getbuffer()) 
            zipf.write(doc1_path, os.path.basename(doc1_path))  
        if doc2 is not None:
            contador = 1
            for uploaded_file in doc2:
                file_type = str(uploaded_file.name).split(".")[-1]
                doc2_path = os.path.join(output_dir,  f"DocumentoAdicional_{contador}.{file_type}") 
                with open(doc2_path, "wb") as f:
                    f.write(uploaded_file.getbuffer()) 
                zipf.write(doc2_path, os.path.basename(doc2_path))  
                contador += 1

    os.remove(archivo_docx)
    if doc1 is not None:
        os.remove(doc1_path)
    if doc2 is not None:
        os.remove(doc2_path)
    
    
    print(f'Documento y archivos añadidos al ZIP: {nombre_zip}')

    return documento, archivo_zip

def crear_documento_consulting_services(data):
    documento = Document()
    
    # Agregar el título
    titulo = documento.add_paragraph()
    run_titulo = titulo.add_run('Consulting Services')
    run_titulo.font.size = Pt(16)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def agregar_encabezado(texto):
        parrafo = documento.add_paragraph()
        run = parrafo.add_run(texto)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Azul
        parrafo.space_after = Pt(6)

    def agregar_bullet_point(campo, valor):
        parrafo = documento.add_paragraph(style='List Bullet')
        run_campo = parrafo.add_run(f"{campo}: ")
        run_campo.font.size = Pt(11)
        run_campo.font.bold = True
        run_valor = parrafo.add_run(f"{valor}")
        run_valor.font.size = Pt(11)

    # Agregar secciones
    agregar_encabezado("Detalle de la Actividad")
    agregar_bullet_point("Nombre", data.get("nombre_necesidades_cs", ""))
    agregar_bullet_point("Owner", data.get('owner_cs', ''))
    if data.get("delegate_cs", "") != "":
        agregar_bullet_point("Delegate", data.get('delegate_cs', ''))
    else:
        agregar_bullet_point("Delegate", "N/A")
    agregar_bullet_point("Fecha de inicio", data.get("start_date_cs", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Fecha de fin", data.get("end_date_cs", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Presupuesto estimado", f"{data.get('presupuesto_estimado_cs', 0)} €")
    if data.get("producto_asociado_cs", "") != "":
        agregar_bullet_point("Producto asociado", data.get("producto_asociado_cs", ""))
    else:
        agregar_bullet_point("Producto asociado", "N/A")
    agregar_bullet_point("Estado de aprobación", data.get("estado_aprobacion_cs", "N/A"))
    agregar_bullet_point("Necesidad de la reunión", data.get("necesidad_reunion_cs", ""))
    agregar_bullet_point("Descripción del servicio", data.get("descripcion_servicio_cs", ""))

    agregar_encabezado("Detalle nº Consultores")
    agregar_bullet_point("Número de consultores", data.get("numero_consultores_cs", ""))
    criterios = ", ".join(data.get("criterios_seleccion_cs", []))
    agregar_bullet_point("Criterios del destinatario", criterios)
    if data.get("justificacion_numero_participantes_cs", "") != "":
        agregar_bullet_point("Justificación", data.get("justificacion_numero_participantes_cs", ""))
    else:
        agregar_bullet_point("Justificación", "N/A")
    
    agregar_encabezado("Detalles de los Consultores")
    tabla = documento.add_table(rows=1, cols=8)
    tabla.style = 'Table Grid'
    encabezados = ["Nombre", "DNI", "Tier", "Centro de trabajo", "Email", "Cobra a través de sociedad", "Honorarios", "Tiempos"]
    hdr_cells = tabla.rows[0].cells
    for i, encabezado in enumerate(encabezados):
        hdr_cells[i].text = encabezado
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    for participante in data.get("participantes_cs", {}).values():
        id_participante = participante["id"]
        row_cells = tabla.add_row().cells
        
        # Fill in basic fields
        row_cells[0].text = participante.get(f"nombre_{id_participante}", "").get("result", "").rsplit('-', 1)[0] 
        row_cells[1].text = participante.get(f"dni_{id_participante}", "")
        row_cells[2].text = participante.get(f"tier_{id_participante}", "")
        row_cells[3].text = participante.get(f"centro_trabajo_{id_participante}", "")
        row_cells[4].text = participante.get(f"email_{id_participante}", "")
        if participante.get(f"cobra_sociedad_{id_participante}", "") == "Sí":
            row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") + ", con sociedad: " + participante.get(f"nombre_sociedad_{id_participante}", "")
        else:
            row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") 
        honorarios = participante.get(f"honorarios_{id_participante}", 0)
        row_cells[6].text = f"{honorarios} €"

        # Add preparation and presentation times
        prep_horas = participante.get(f"preparacion_horas_{id_participante}", 0)
        prep_mins = participante.get(f"preparacion_minutos_{id_participante}", 0)
        pon_horas = participante.get(f"ponencia_horas_{id_participante}", 0)
        pon_mins = participante.get(f"ponencia_minutos_{id_participante}", 0)
        
        row_cells[7].text = f"Preparación: {prep_horas}horas y {prep_mins}minutos, Ponencia: {pon_horas}horas y {pon_mins}minutos"

    nombre_zip = f"Consulting_Services {data.get('nombre_necesidades_cs', '')}.zip"
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    archivo_zip = os.path.join(output_dir, nombre_zip)

    with zipfile.ZipFile(archivo_zip, 'w') as zipf:
        nombre_archivo = 'Consulting_Services_Plantilla.docx'
        archivo_docx = os.path.join(output_dir, nombre_archivo)
        documento.save(archivo_docx)
        zipf.write(archivo_docx, os.path.basename(archivo_docx))

        doc1 = data.get("documentosubido_1_cs", None)
        doc2 = data.get("documentosubido_2_cs", None)

        if doc1 is not None:
            file_type = str(doc1.name).split(".")[-1]
            doc1_path = os.path.join(output_dir, f"AgendaEvento.{file_type}")  
            with open(doc1_path, "wb") as f:
                f.write(doc1.getbuffer())  
            zipf.write(doc1_path, os.path.basename(doc1_path)) 
        
        if doc2 is not None:
            contador = 1
            for uploaded_file in doc2:
                file_type = str(uploaded_file.name).split(".")[-1]
                doc2_path = os.path.join(output_dir,  f"DocumentoAdicional_{contador}.{file_type}") 
                with open(doc2_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                zipf.write(doc2_path, os.path.basename(doc2_path))
                contador += 1



    os.remove(archivo_docx)
    if doc1 is not None:
        os.remove(doc1_path)
    if doc2 is not None:
        os.remove(doc2_path)

    print(f'Documento y archivos añadidos al ZIP: {nombre_zip}')

    return documento, archivo_zip


def crear_documento_speaking(data):
    documento = Document()

    # Agregar el título
    titulo = documento.add_paragraph()
    run_titulo = titulo.add_run('Speaking Services')
    run_titulo.font.size = Pt(16)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def agregar_encabezado(texto):
        parrafo = documento.add_paragraph()
        run = parrafo.add_run(texto)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Azul
        parrafo.space_after = Pt(6)

    def agregar_bullet_point(campo, valor):
        parrafo = documento.add_paragraph(style='List Bullet')
        run_campo = parrafo.add_run(f"{campo}: ")
        run_campo.font.size = Pt(11)
        run_campo.font.bold = True
        run_valor = parrafo.add_run(f"{valor}")
        run_valor.font.size = Pt(11)

    # Agregar secciones
    agregar_encabezado("Detalles del Evento")
    agregar_bullet_point("Nombre", data.get("nombre_evento_ss", ""))
    agregar_bullet_point("Owner", data.get('owner_ss', ''))
    if data.get("delegate_ss", "") != "":
        agregar_bullet_point("Delegate", data.get('delegate_ss', ''))
    else:
        agregar_bullet_point("Delegate", "N/A")
    agregar_bullet_point("Descripción", data.get("descripcion_objetivo_ss", ""))
    agregar_bullet_point("Fecha de inicio", data.get("start_date_ss", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Fecha de fin", data.get("end_date_ss", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Nº asistentes totales", data.get("num_asistentes_totales_ss", ""))
    agregar_bullet_point("Tipo de evento", data.get("tipo_evento_ss", ""))
    if data.get("sede_ss") == "":
        agregar_bullet_point("Sede", "N/A")
    else:
        agregar_bullet_point("Sede", data.get("sede_ss", "N/A"))

    if data.get("ciudad_ss") == "":
        agregar_bullet_point("Ciudad", "N/A")
    else:
        agregar_bullet_point("Ciudad", data.get("ciudad_ss", "N/A"))

    agregar_bullet_point("Público objetivo del programa", data.get("publico_objetivo_ss", ""))

    agregar_encabezado("Detalles de la Actividad")
    agregar_bullet_point("Presupuesto total estimado",  f"{data.get('presupuesto_estimado_ss', 0)} €")
    if data.get("producto_asociado_ss", "") != "":
        agregar_bullet_point("Producto asociado", data.get("producto_asociado_ss", ""))
    else:
        agregar_bullet_point("Producto asociado", "N/A")
    agregar_bullet_point("Necesidad de la reunión y resultados esperados", data.get("necesidad_reunion_ss", ""))
    agregar_bullet_point("Descripción del servicio", data.get("servicio_ss", ""))
    
    agregar_encabezado("Detalle nº de Ponentes")
    agregar_bullet_point("Nº de ponentes", data.get("num_ponentes_ss", ""))
    criterios = ", ".join(data.get("criterios_seleccion_ss", []))
    agregar_bullet_point("Criterios de selección", criterios)

    agregar_encabezado("Logística de los Ponentes")
    agregar_bullet_point("Desplazamiento de ponentes", data.get("desplazamiento_ponentes_ss", ""))
    agregar_bullet_point("Alojamiento de ponentes", data.get("alojamiento_ponentes_ss", ""))
    if data.get("alojamiento_ponentes_ss", "") == "Sí":
        agregar_bullet_point("Nº de noches", data.get("num_noches_ss", ""))
        agregar_bullet_point("Hotel", data.get("hotel_ss", ""))
    

    agregar_encabezado("Detalles de los Ponentes")
    def agregar_tabla_participantes(participantes):
        tabla = documento.add_table(rows=1, cols=8)
        tabla.style = 'Table Grid'  # Aplicar bordes a la tabla
        encabezados = ["Nombre y Apellidos", "DNI", "Tier", "Centro de trabajo", "Email", "Cobra a través de sociedad", "Honorarios", "Tiempos"]
        hdr_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            hdr_cells[i].text = encabezado
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i]._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w')))
        
        for participante in participantes.values():
            id_participante = participante["id"]
            row_cells = tabla.add_row().cells
            # row_cells[0].text = participante.get(f"nombre_{id_participante}", "").split('-')[0]
            row_cells[0].text = participante.get(f"nombre_{id_participante}", "").get("result", "").rsplit('-', 1)[0] 
            #row_cells[0].text = "-".join(participante.get(f"nombre_{id_participante}", "").split('-')[:-1])
            row_cells[1].text = participante.get(f"dni_{id_participante}", "")
            row_cells[2].text = participante.get(f"tier_{id_participante}", "")
            row_cells[3].text = participante.get(f"centro_trabajo_{id_participante}", "")
            row_cells[4].text = participante.get(f"email_{id_participante}", "")
            if participante.get(f"cobra_sociedad_{id_participante}", "") == "Sí":
                row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") + ", con sociedad: " + participante.get(f"nombre_sociedad_{id_participante}", "")
            else:
                row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") 
            row_cells[6].text = str(participante.get(f"honorarios_{id_participante}", ""))
            row_cells[7].text = f"Preparación: {participante.get(f'preparacion_horas_{id_participante}', '')} horas y {participante.get(f'preparacion_minutos_{id_participante}', '')} minutos, Ponencia: {participante.get(f'ponencia_horas_{id_participante}', '')} horas y {participante.get(f'ponencia_minutos_{id_participante}', '')} minutos"
            
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w')))
    
    agregar_tabla_participantes(data.get("participantes_ss", {}))
    
    nombre_zip = f"Speaking_Service_Merck_Program {data.get('nombre_evento_ss', '')}.zip"
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    archivo_zip = os.path.join(output_dir, nombre_zip)

    with zipfile.ZipFile(archivo_zip, 'w') as zipf:
        nombre_archivo = 'Speaking_Services_Plantilla.docx'
        archivo_docx = os.path.join(output_dir, nombre_archivo)
        documento.save(archivo_docx)
        zipf.write(archivo_docx, os.path.basename(archivo_docx))

        doc1 = data.get("documentosubido_1_ss", None)
        doc2 = data.get("documentosubido_2_ss", None)
        doc3 = data.get("documentosubido_3_ss", None)

        if doc1 is not None:
            file_type = str(doc1.name).split(".")[-1]
            doc1_path = os.path.join(output_dir, f"AgendaEvento.{file_type}")  
            with open(doc1_path, "wb") as f:
                f.write(doc1.getbuffer()) 
            zipf.write(doc1_path, os.path.basename(doc1_path))  

        if doc2 is not None:
            file_type = str(doc2.name).split(".")[-1]
            doc2_path = os.path.join(output_dir, f"Contratos.{file_type}")  
            with open(doc2_path, "wb") as f:
                f.write(doc2.getbuffer()) 
            zipf.write(doc2_path, os.path.basename(doc2_path))  

        if doc3 is not None:
            contador = 1
            for uploaded_file in doc3:
                file_type = str(uploaded_file.name).split(".")[-1]
                doc3_path = os.path.join(output_dir, f"DocumentoAdicional_{contador}.{file_type}")  
                with open(doc3_path, "wb") as f:
                    f.write(uploaded_file.getbuffer()) 
                zipf.write(doc3_path, os.path.basename(doc3_path))  
                contador += 1


    os.remove(archivo_docx)
    if doc1 is not None:
        os.remove(doc1_path)
    if doc2 is not None:
        os.remove(doc2_path)
    if doc3 is not None:
        os.remove(doc3_path)

    print(f'Documento y archivos añadidos al ZIP: {nombre_zip}')

    return documento, archivo_zip

    

def crear_documento_speaking_reducido(data):
    documento = Document()

    # Agregar el título
    titulo = documento.add_paragraph()
    run_titulo = titulo.add_run('Speaking Engagement Paragüas')
    run_titulo.font.size = Pt(16)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 0, 128)  # Azul oscuro
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def agregar_encabezado(texto):
        parrafo = documento.add_paragraph()
        run = parrafo.add_run(texto)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Azul
        parrafo.space_after = Pt(6)

    def agregar_bullet_point(campo, valor):
        parrafo = documento.add_paragraph(style='List Bullet')
        run_campo = parrafo.add_run(f"{campo}: ")
        run_campo.font.size = Pt(11)
        run_campo.font.bold = True
        run_valor = parrafo.add_run(f"{valor}")
        run_valor.font.size = Pt(11)

    # Agregar secciones
    agregar_encabezado("Detalles del Evento")
    agregar_bullet_point("Nombre", data.get("nombre_evento_ss", ""))
    agregar_bullet_point("Owner", data.get('owner_ss', ''))
    if data.get("delegate_ss", "") != "":
        agregar_bullet_point("Delegate", data.get('delegate_ss', ''))
    else:
        agregar_bullet_point("Delegate", "N/A")
    agregar_bullet_point("Fecha de inicio", data.get("start_date_ss", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Fecha de fin", data.get("end_date_ss", "").strftime("%d/%m/%Y"))
    agregar_bullet_point("Tipo de evento", data.get("tipo_evento_ss", ""))
    if data.get("sede_ss") == "":
        agregar_bullet_point("Sede", "N/A")
    else:
        agregar_bullet_point("Sede", data.get("sede_ss", "N/A"))

    if data.get("ciudad_ss") == "":
        agregar_bullet_point("Ciudad", "N/A")
    else:
        agregar_bullet_point("Ciudad", data.get("ciudad_ss", "N/A"))
    
    agregar_encabezado("Detalles de los Ponentes")
    def agregar_tabla_participantes(participantes):
        tabla = documento.add_table(rows=1, cols=8)
        tabla.style = 'Table Grid'  # Aplicar bordes a la tabla
        encabezados = ["Nombre y Apellidos", "DNI", "Tier", "Centro de trabajo", "Email", "Cobra a través de sociedad", "Honorarios", "Tiempos"]
        hdr_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            hdr_cells[i].text = encabezado
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i]._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w')))
        
        for participante in participantes.values():
            id_participante = participante["id"]
            row_cells = tabla.add_row().cells
            row_cells[0].text = participante.get(f"nombre_{id_participante}", "").get("result", "").rsplit('-', 1)[0] 
            row_cells[1].text = participante.get(f"dni_{id_participante}", "")
            row_cells[2].text = participante.get(f"tier_{id_participante}", "")
            row_cells[3].text = participante.get(f"centro_trabajo_{id_participante}", "")
            row_cells[4].text = participante.get(f"email_{id_participante}", "")
            if participante.get(f"cobra_sociedad_{id_participante}", "") == "Sí":
                row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") + ", con sociedad: " + participante.get(f"nombre_sociedad_{id_participante}", "")
            else:
                row_cells[5].text = participante.get(f"cobra_sociedad_{id_participante}", "") 
            row_cells[6].text = str(participante.get(f"honorarios_{id_participante}", ""))
            row_cells[7].text = f"Preparación: {participante.get(f'preparacion_horas_{id_participante}', '')} horas y {participante.get(f'preparacion_minutos_{id_participante}', '')} minutos, Ponencia: {participante.get(f'ponencia_horas_{id_participante}', '')} horas y {participante.get(f'ponencia_minutos_{id_participante}', '')} minutos"
            
            for cell in row_cells:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w')))
    
    agregar_tabla_participantes(data.get("participantes_ss", {}))
    
    nombre_zip = f"Speaking_Service_Paragüas {data.get('nombre_evento_ss', '')}.zip"
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    archivo_zip = os.path.join(output_dir, nombre_zip)

    with zipfile.ZipFile(archivo_zip, 'w') as zipf:
        nombre_archivo = 'Speaking_Services_Paragüas_Plantilla.docx'
        archivo_docx = os.path.join(output_dir, nombre_archivo)
        documento.save(archivo_docx)
        zipf.write(archivo_docx, os.path.basename(archivo_docx))
        os.remove(archivo_docx)  
    
    print(f'Documento y archivos añadidos al ZIP: {nombre_zip}')

    return documento, archivo_zip
